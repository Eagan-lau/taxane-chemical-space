from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image


PLACEHOLDER_PATTERN = re.compile(r"^\{\{([A-Z0-9_]+)\}\}$")


def _set_cell_borderless(paragraph) -> None:
    paragraph.paragraph_format.space_after = Pt(0)


def _set_font(run, name: str = "Times New Roman", size: float = 11, bold: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_inline_markdown(paragraph, text: str, size: float = 11) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        run = paragraph.add_run(value)
        _set_font(run, size=size, bold=bold)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    heading_sizes = {1: 15, 2: 13, 3: 11.5}
    for level, size in heading_sizes.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Taxane reaction-grammar study")
    _set_font(run, size=8)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _markdown_blocks(markdown_text: str) -> list[str]:
    """Join source-wrapped prose while preserving structural Markdown lines."""
    blocks: list[str] = []
    paragraph_lines: list[str] = []
    list_prefix: str | None = None

    def flush_paragraph() -> None:
        nonlocal list_prefix
        if paragraph_lines:
            text = " ".join(part.strip() for part in paragraph_lines)
            blocks.append(f"{list_prefix} {text}" if list_prefix else text)
            paragraph_lines.clear()
        list_prefix = None

    for raw_line in markdown_text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            flush_paragraph()
            continue
        if stripped.startswith("#") or bool(PLACEHOLDER_PATTERN.match(stripped)):
            flush_paragraph()
            blocks.append(stripped)
            continue
        bullet_match = re.match(r"^(-|\d+\.)\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            list_prefix = bullet_match.group(1)
            paragraph_lines.append(bullet_match.group(2))
            continue
        paragraph_lines.append(stripped)
        if raw_line.endswith("  "):
            flush_paragraph()
    flush_paragraph()
    return blocks


def markdown_to_docx(
    markdown_text: str,
    output_path: Path,
    image_map: dict[str, Path],
    title: str,
    caption_map: dict[str, str] | None = None,
) -> None:
    caption_map = caption_map or {}
    document = Document()
    _configure_document(document)
    document.core_properties.title = title
    document.core_properties.subject = "Taxane reaction-grammar chemical-space analysis"
    document.core_properties.keywords = "taxane, reaction SMARTS, reaction grammar, chemical space"

    footer = document.sections[0].footer.paragraphs[0]
    footer.add_run(" | ")
    _add_page_number(footer)

    for line in _markdown_blocks(markdown_text):
        placeholder = PLACEHOLDER_PATTERN.match(line)
        if placeholder:
            key = placeholder.group(1)
            if key not in image_map:
                raise KeyError(f"Unresolved image placeholder: {key}")
            document.add_page_break()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            with Image.open(image_map[key]) as image:
                aspect_ratio = image.width / image.height
            max_width = 6.5
            max_height = 7.2
            if max_width / aspect_ratio <= max_height:
                paragraph.add_run().add_picture(
                    str(image_map[key]),
                    width=Inches(max_width),
                )
            else:
                paragraph.add_run().add_picture(
                    str(image_map[key]),
                    height=Inches(max_height),
                )
            caption = caption_map.get(key, "").strip()
            if caption:
                caption_blocks = _markdown_blocks(caption)
                for block_index, block in enumerate(caption_blocks):
                    caption_paragraph = document.add_paragraph()
                    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    caption_paragraph.paragraph_format.space_before = Pt(
                        1 if block_index else 0
                    )
                    caption_paragraph.paragraph_format.space_after = Pt(2)
                    _add_inline_markdown(caption_paragraph, block, size=9)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(14)
            _add_inline_markdown(paragraph, line[2:], size=18)
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("#### "):
            document.add_heading(line[5:], level=3)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_inline_markdown(paragraph, line[2:])
            continue
        if re.match(r"^\d+\. ", line):
            paragraph = document.add_paragraph(style="List Number")
            _add_inline_markdown(paragraph, re.sub(r"^\d+\. ", "", line))
            continue
        paragraph = document.add_paragraph()
        _add_inline_markdown(paragraph, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    runtime = output_dir / ".libreoffice_runtime"
    profile = output_dir / ".libreoffice_profile"
    runtime.mkdir(exist_ok=True)
    profile.mkdir(exist_ok=True)
    command = [
        shutil.which("libreoffice") or "libreoffice",
        "--headless",
        f"-env:UserInstallation=file://{profile.resolve()}",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_dir),
        str(docx_path),
    ]
    result = subprocess.run(
        command,
        check=False,
        capture_output=True,
        text=True,
        env={
            "HOME": str(profile),
            "XDG_RUNTIME_DIR": str(runtime),
            "PATH": str(Path(command[0]).parent) + ":/usr/bin:/bin",
        },
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed for {docx_path}\n"
            f"stdout:\n{result.stdout}\nstderr:\n{result.stderr}"
        )
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.is_file() or pdf_path.stat().st_size == 0:
        raise RuntimeError(
            f"LibreOffice did not produce {pdf_path}\n"
            f"stdout:\n{result.stdout}\nstderr:\n{result.stderr}"
        )
    shutil.rmtree(runtime, ignore_errors=True)
    shutil.rmtree(profile, ignore_errors=True)
    return pdf_path
